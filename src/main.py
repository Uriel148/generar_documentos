from docx import Document

def reemplazar_palabra_en_docx(documento_entrada,cambio_palabras,nombre):
    # Cargar el documento
    doc = Document(documento_entrada)
    
    # Función para reemplazar en elementos principales
    def reemplazar_en_elemento(elemento):
        if hasattr(elemento, 'runs'):
            for run in elemento.runs:
                for palabra_buscar, palabra_reemplazo in cambio_palabras.items():
                    if palabra_buscar in run.text:
                        run.text = run.text.replace(palabra_buscar, palabra_reemplazo)
    
    # Procesar párrafos principales
    for párrafo in doc.paragraphs:
        reemplazar_en_elemento(párrafo)
    
    # Procesar tablas
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for párrafo in celda.paragraphs:
                    reemplazar_en_elemento(párrafo)
    
    # Procesar encabezados y pies de página
    for sección in doc.sections:
        for encabezado in [sección.header, sección.first_page_header]:
            if encabezado:
                for párrafo in encabezado.paragraphs:
                    reemplazar_en_elemento(párrafo)
        
        for pie in [sección.footer, sección.first_page_footer]:
            if pie:
                for párrafo in pie.paragraphs:
                    reemplazar_en_elemento(párrafo)
    
    # Guardar el documento modificado
    documento_salida = documento_entrada.split(".")[0] + "_" + nombre + ".docx"
    doc.save(documento_salida)
    print(f"\nDocumento modificado guardado como: {documento_salida}")

if __name__ == "__main__":
    documento_entrada = r"doc\test2.docx"
    palabras_buscar = ["XXXXX","YYYYY","ZZZZZ"] 
    palabras_reemplazo = ["exitosa_1","exitosa_2","exitosa_3"]
    cambios_palabras = [{"XXXXX": "exitosa_1","YYYYY":"exitosa_2","ZZZZZ":"exitosa_3"},
                        {"XXXXX": "test_1","YYYYY":"test_2","ZZZZZ":"test_3"},
                        {"XXXXX": "ejemplo_1","YYYYY":"ejemplo_2","ZZZZZ":"ejemplo_3"}] 
    nombres = ["exitosa","test","ejemplo"] 
    
    #Falta probar con tablas y con imagenes
    for cambio_palabras, nombre in zip(cambios_palabras, nombres):
        reemplazar_palabra_en_docx(documento_entrada,cambio_palabras,nombre)
    